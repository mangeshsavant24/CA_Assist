"""
agents/document_agent.py
─────────────────────────────────────────────────────────────────
Robust, LLM-independent document extraction pipeline.

Tier 1 (always runs):  Tesseract OCR + Regex heuristics
Tier 2 (optional):     LLM enhancement when Ollama / OpenAI is available

Public API:
  process_document(file_path, filename) -> dict    ← new master function
  DocumentAgent().handle(file_path)   -> dict      ← backward-compat wrapper
"""

import json
import logging
import os
import re
import unicodedata
import warnings
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# PART 1 — TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════

def _clean_text(text: str) -> str:
    """Normalize and clean raw extracted text."""
    if not text:
        return ""
    # Unicode normalization (handles ligatures, fancy quotes, etc.)
    text = unicodedata.normalize("NFKD", text)
    # Strip null bytes
    text = text.replace("\x00", "")
    # Replace form-feed with newline
    text = text.replace("\f", "\n")
    # Collapse 3+ consecutive spaces/tabs to a single space
    text = re.sub(r"[ \t]{3,}", " ", text)
    # Collapse 3+ consecutive newlines to two
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _preprocess_image_for_ocr(image):
    """
    Prepare a PIL image for Tesseract OCR:
      - Convert to grayscale
      - Scale up to at least ~300 DPI equivalent if the image is small
      - Apply sharpening to improve recognition
    """
    from PIL import Image, ImageFilter

    # Grayscale
    if image.mode != "L":
        image = image.convert("L")

    # Scale up small images (assume 72 DPI source → 300 DPI target)
    w, h = image.size
    if w < 1000 or h < 1000:
        scale = max(300 / 72, 1.5)
        image = image.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    # Sharpen
    image = image.filter(ImageFilter.SHARPEN)
    return image


def _run_tesseract(image) -> str:
    """Run Tesseract with optimal config on a prepared PIL image."""
    import pytesseract
    return pytesseract.image_to_string(image, config="--oem 3 --psm 6")


# ───────────────────────────── PDF ────────────────────────────────

def _extract_pdf_text(file_path: str) -> str:
    """
    Extract text from a PDF.
    Order: pdfplumber → PyPDF2 → pdf2image+Tesseract (scanned PDFs).
    """
    text = ""

    # 1. Try pdfplumber (better layout handling)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            parts = [p.extract_text() or "" for p in pdf.pages]
        text = "\n\n".join(p for p in parts if p.strip())
    except ImportError:
        logger.debug("pdfplumber not installed, trying PyPDF2")
    except Exception as exc:
        logger.warning("pdfplumber failed on %s: %s", file_path, exc)

    # 2. Fallback to PyPDF2
    if len(text.strip()) < 50:
        try:
            import PyPDF2
            with open(file_path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                parts = []
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        parts.append(extracted)
            text = "\n".join(parts)
        except ImportError:
            logger.debug("PyPDF2 not installed")
        except Exception as exc:
            logger.warning("PyPDF2 failed on %s: %s", file_path, exc)

    # 3. If still sparse → OCR via pdf2image (handles scanned PDFs)
    if len(text.strip()) < 50:
        try:
            from pdf2image import convert_from_path
            pages = convert_from_path(file_path, dpi=300)
            ocr_parts = []
            for img in pages:
                processed = _preprocess_image_for_ocr(img)
                t = _run_tesseract(processed)
                if t.strip():
                    ocr_parts.append(t)
            text = "\n\n".join(ocr_parts)
            logger.info("Used OCR fallback for PDF: %d chars extracted", len(text))
        except ImportError:
            logger.warning(
                "pdf2image / pytesseract not installed — cannot OCR scanned PDF. "
                "Install: pip install pdf2image pytesseract  "
                "(also needs poppler and Tesseract binaries)"
            )
        except Exception as exc:
            logger.warning("PDF OCR fallback failed: %s", exc)

    return text


# ─────────────────────────── Images ──────────────────────────────

def _extract_image_text(file_path: str) -> str:
    """Extract text from an image via Tesseract OCR with preprocessing."""
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(file_path)
        processed = _preprocess_image_for_ocr(image)
        return _run_tesseract(processed)

    except ImportError:
        logger.error(
            "PIL (Pillow) or pytesseract not installed. "
            "Install with: pip install Pillow pytesseract"
        )
        return ""
    except Exception as exc:
        # Check specifically for Tesseract binary missing
        if "tesseract" in str(exc).lower() or "not found" in str(exc).lower():
            logger.error(
                "Tesseract OCR binary not found. "
                "Download from: https://github.com/UB-Mannheim/tesseract/wiki (Windows). "
                "Then add its install directory to the system PATH."
            )
        else:
            logger.warning("Image OCR failed for %s: %s", file_path, exc)
        return ""


# ─────────────────────────── DOCX ────────────────────────────────

def _extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file (paragraphs + table cells)."""
    try:
        import docx
        doc = docx.Document(file_path)
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        logger.warning("python-docx not installed — cannot extract .docx files")
        return ""
    except Exception as exc:
        logger.warning("DOCX extraction failed for %s: %s", file_path, exc)
        return ""


# ─────────────────────────── Master ──────────────────────────────

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from any supported document file.

    Supported: PDF, JPG, JPEG, PNG, WEBP, BMP, TIFF, DOCX, TXT
    Returns cleaned text string (may be empty on failure).
    """
    if not file_path or not os.path.exists(file_path):
        logger.warning("File does not exist: %s", file_path)
        return ""

    ext = Path(file_path).suffix.lower().lstrip(".")

    if ext == "pdf":
        text = _extract_pdf_text(file_path)
    elif ext in ("jpg", "jpeg", "png", "webp", "bmp", "tiff", "tif"):
        text = _extract_image_text(file_path)
    elif ext in ("docx", "doc"):
        text = _extract_docx_text(file_path)
    elif ext == "txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
                text = fh.read()
        except Exception as exc:
            logger.warning("TXT read failed for %s: %s", file_path, exc)
            text = ""
    else:
        logger.warning("Unsupported file type: .%s", ext)
        text = ""

    return _clean_text(text)


# ═══════════════════════════════════════════════════════════════════
# PART 2 — HEURISTIC DOCUMENT CLASSIFIER
# ═══════════════════════════════════════════════════════════════════

# Keyword rules per document type.
# text_threshold = minimum text keyword matches required.
_DOCUMENT_TYPE_RULES: dict = {
    "salary_slip": {
        "filename_keywords": ["salary", "payslip", "pay slip", "pay-slip", "payroll", "ctc"],
        "text_keywords": [
            "basic salary", "basic pay", "gross salary", "gross pay",
            "net salary", "net pay", "take home", "hra", "house rent",
            "pf", "provident fund", "esic", "tds", "deductions",
            "earnings", "allowances", "ctc", "cost to company",
            "employee code", "employee id", "pay period",
        ],
        "text_threshold": 3,
    },
    "form_16": {
        "filename_keywords": ["form16", "form 16", "form-16"],
        "text_keywords": [
            "certificate under section 203", "tds certificate",
            "form no. 16", "assessment year", "pan of deductor",
            "tan of deductor", "acknowledgement number",
            "salary as per provisions",
        ],
        "text_threshold": 3,
    },
    "itr_document": {
        "filename_keywords": ["itr", "income tax return"],
        "text_keywords": [
            "acknowledgement number", "assessment year",
            "income tax return", "itr-1", "itr-2", "itr-3",
            "gross total income", "deductions under chapter vi",
            "tax payable", "refund",
        ],
        "text_threshold": 3,
    },
    "invoice": {
        "filename_keywords": ["invoice", "bill", "receipt"],
        "text_keywords": [
            "invoice no", "invoice number", "bill to", "ship to",
            "gstin", "gst number", "hsn", "sac code",
            "taxable value", "cgst", "sgst", "igst",
            "total amount", "subtotal", "tax invoice",
        ],
        "text_threshold": 3,
    },
    "gst_return": {
        "filename_keywords": ["gstr", "gst return"],
        "text_keywords": [
            "gstr-1", "gstr-3b", "gstr-2a", "gstr-9",
            "outward supplies", "inward supplies",
            "input tax credit", "itc",
            "table 3", "table 4", "gstin",
            "filing period", "taxpayer name",
        ],
        "text_threshold": 3,
    },
    "forex_document": {
        "filename_keywords": [
            "forex", "foreign exchange", "fema", "swift",
            "wire transfer", "remittance", "currency",
        ],
        "text_keywords": [
            "exchange rate", "foreign currency", "usd", "eur", "gbp",
            "remittance", "wire transfer", "swift code", "iban",
            "nostro", "vostro", "spot rate", "forward rate",
            "forex", "fema", "rbi", "ad category",
            "currency exposure", "hedging",
        ],
        "text_threshold": 3,
    },
    "fund_document": {
        "filename_keywords": ["nav", "fund", "portfolio", "mutual fund"],
        "text_keywords": [
            "net asset value", "nav", "aum",
            "assets under management", "portfolio",
            "scheme name", "folio number", "units",
            "dividend", "growth", "redemption",
            "fund manager", "benchmark",
        ],
        "text_threshold": 3,
    },
    "tax_notice": {
        "filename_keywords": [],
        "text_keywords": [
            "notice under section", "demand notice",
            "income tax notice", "compliance notice",
            "response required", "scrutiny",
            "assessment order",
        ],
        "text_threshold": 2,
    },
}

_REGIME_RELEVANT: frozenset = frozenset({"salary_slip", "form_16", "itr_document"})
_FOREX_RELEVANT: frozenset = frozenset({"forex_document"})
_FUND_RELEVANT: frozenset = frozenset({"fund_document"})
_ALL_KNOWN_TYPES: frozenset = frozenset(_DOCUMENT_TYPE_RULES) | {"unknown"}


def classify_document_heuristic(text: str, filename: str) -> dict:
    """
    Classify a document using filename keywords and text keyword matching.

    Scoring:
      filename match  = +5 points (strong signal)
      each text match = +1 point

    Confidence:
      filename + text = 0.70 – 0.95
      filename only   = 0.75
      text only       = 0.50 – 0.85
      no match        = 0.30 (unknown)

    Returns:
      {
        "document_type": str,
        "confidence": float,
        "is_relevant_for_regime": bool,
        "is_relevant_for_forex": bool,
        "is_relevant_for_fund": bool,
        "detected_by": "heuristic"
      }
    """
    filename_lower = (filename or "").lower()
    text_lower = (text or "").lower()

    best_type = "unknown"
    best_score = -1
    best_confidence = 0.3

    for doc_type, rules in _DOCUMENT_TYPE_RULES.items():
        filename_matched = any(kw in filename_lower for kw in rules["filename_keywords"])
        text_matches = sum(1 for kw in rules["text_keywords"] if kw in text_lower)
        threshold = rules["text_threshold"]
        text_matched = text_matches >= threshold

        if not filename_matched and not text_matched:
            continue

        score = text_matches + (5 if filename_matched else 0)
        total_keywords = max(len(rules["text_keywords"]), 1)

        if filename_matched and text_matched:
            confidence = min(0.95, 0.70 + (text_matches / total_keywords) * 0.25)
        elif filename_matched:
            confidence = 0.75
        else:
            confidence = min(0.85, 0.50 + (text_matches / total_keywords) * 0.35)

        if score > best_score:
            best_score = score
            best_type = doc_type
            best_confidence = confidence

    return {
        "document_type": best_type,
        "confidence": round(best_confidence, 3),
        "is_relevant_for_regime": best_type in _REGIME_RELEVANT,
        "is_relevant_for_forex": best_type in _FOREX_RELEVANT,
        "is_relevant_for_fund": best_type in _FUND_RELEVANT,
        "detected_by": "heuristic",
    }


# ═══════════════════════════════════════════════════════════════════
# PART 3 — FIELD EXTRACTION BY DOCUMENT TYPE
# ═══════════════════════════════════════════════════════════════════

def _parse_amount(raw: Optional[str]) -> Optional[float]:
    """Convert '₹1,23,456' or '1,23,456.00' to float. Returns None on failure."""
    if not raw:
        return None
    cleaned = re.sub(r"[₹,\s]", "", str(raw))
    try:
        return float(cleaned)
    except (ValueError, TypeError):
        return None


def _first_match(text: str, patterns: list) -> Optional[str]:
    """
    Try each pattern against text (IGNORECASE).
    Return the value of the last capturing group of the first match found.
    """
    for pattern in patterns:
        m = re.search(pattern, text, re.IGNORECASE)
        if m and m.lastindex:
            return m.group(m.lastindex)
    return None


def _extract_date(text: str) -> Optional[str]:
    """Extract the first recognisable date string from text."""
    patterns = [
        r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{4})\b",                           # DD/MM/YYYY or DD-MM-YYYY
        r"\b(\d{4}-\d{2}-\d{2})\b",                                      # YYYY-MM-DD (ISO)
        r"\b(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4})\b",  # DD Mon YYYY
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1)
    return None


# ──────────────────────── Salary Slip ────────────────────────────

def _extract_salary_slip_fields(text: str) -> dict:
    fields: dict = {}

    fields["gross_salary"] = _parse_amount(_first_match(text, [
        r"gross\s*(?:salary|pay|earnings|income)\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"gross\s*ctc\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"total\s*earnings\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    fields["basic_pay"] = _parse_amount(_first_match(text, [
        r"basic\s*(?:pay|salary)\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    fields["hra"] = _parse_amount(_first_match(text, [
        r"h\.?r\.?a\.?\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"house\s*rent\s*allowance\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    fields["special_allowance"] = _parse_amount(_first_match(text, [
        r"special\s*allowance\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    # EPF first (more specific), then generic PF
    fields["pf_deduction"] = _parse_amount(_first_match(text, [
        r"\bepf\b\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"\bp\.?f\.?\b\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"provident\s*fund\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))
    # Alias for backward compatibility with API schema
    fields["pf"] = fields["pf_deduction"]

    fields["tds_deducted"] = _parse_amount(_first_match(text, [
        r"\btds\b\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"tax\s*deducted\s*(?:at\s*source)?\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"income\s*tax\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    fields["net_pay"] = _parse_amount(_first_match(text, [
        r"net\s*(?:pay|salary|take\s*home)\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"take\s*home\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    fields["professional_tax"] = _parse_amount(_first_match(text, [
        r"professional\s*tax\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"\bp\.?\s*tax\b\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    fields["medical_allowance"] = _parse_amount(_first_match(text, [
        r"medical\s*allowance\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    # Pay period: "March 2025", "Mar 2025", etc.
    m_period = re.search(
        r"\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?"
        r"|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?"
        r"|dec(?:ember)?)\s+(\d{4})\b",
        text,
        re.IGNORECASE,
    )
    if m_period:
        fields["pay_period"] = f"{m_period.group(1).capitalize()} {m_period.group(2)}"

    return fields


# ─────────────────────── Forex Document ──────────────────────────

def _extract_forex_fields(text: str) -> dict:
    fields: dict = {}

    m_pair = re.search(
        r"\b(USD|EUR|GBP|JPY|CHF|AUD|CAD|SGD|AED|SAR)\s*[/]?\s*(INR|USD|EUR|GBP)\b",
        text,
        re.IGNORECASE,
    )
    if m_pair:
        fields["currency_pair"] = f"{m_pair.group(1).upper()}/{m_pair.group(2).upper()}"

    fields["amount"] = _parse_amount(_first_match(text, [
        r"amount\s*[:\-]?\s*[₹$]?\s*([\d,]+\.?\d*)",
        r"principal\s*[:\-]?\s*[₹$]?\s*([\d,]+\.?\d*)",
    ]))

    fields["exchange_rate"] = _parse_amount(_first_match(text, [
        r"exchange\s*rate\s*[:\-]?\s*([\d.]+)",
        r"spot\s*rate\s*[:\-]?\s*([\d.]+)",
        r"\brate\b\s*[:\-]?\s*([\d.]+)",
    ]))

    m_type = re.search(
        r"\b(remittance|wire\s*transfer|forward|spot\s*deal|swap|inward|outward)\b",
        text,
        re.IGNORECASE,
    )
    if m_type:
        fields["transaction_type"] = m_type.group(1).title()

    fields["value_date"] = _extract_date(text)

    return fields


# ─────────────────────── Fund Document ───────────────────────────

def _extract_fund_fields(text: str) -> dict:
    fields: dict = {}

    m_fund = re.search(r"(?:scheme|fund)\s*name\s*[:\-]?\s*([^\n]+)", text, re.IGNORECASE)
    if m_fund:
        fields["fund_name"] = m_fund.group(1).strip()

    fields["nav"] = _parse_amount(_first_match(text, [
        r"\bn\.?a\.?v\.?\b\s*[:\-]?\s*[₹]?\s*([\d.]+)",
    ]))

    fields["units"] = _parse_amount(_first_match(text, [
        r"units\s*[:\-]?\s*([\d,]+\.?\d*)",
    ]))

    fields["total_value"] = _parse_amount(_first_match(text, [
        r"(?:total|current)\s*(?:value|amount)\s*[:\-]?\s*[₹]?\s*([\d,]+)",
    ]))

    m_folio = re.search(r"folio\s*(?:no\.?|number)\s*[:\-]?\s*(\w+)", text, re.IGNORECASE)
    if m_folio:
        fields["folio_number"] = m_folio.group(1)

    return fields


# ──────────────────────── Invoice ────────────────────────────────

def _extract_invoice_fields(text: str) -> dict:
    fields: dict = {}

    m_inv = re.search(
        r"invoice\s*(?:no\.?|#|number)\s*[:\-]?\s*([\w][\w\-/]*)",
        text,
        re.IGNORECASE,
    )
    if m_inv:
        fields["invoice_number"] = m_inv.group(1).strip()

    fields["invoice_date"] = _extract_date(text)

    fields["total_amount"] = _parse_amount(_first_match(text, [
        r"(?:grand\s*total|total\s*amount|total\s*payable|total\s*due)\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
        r"\btotal\b\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)",
    ]))

    m_gstin = re.search(
        r"gstin\s*[:\-]?\s*(\d{2}[A-Z]{5}\d{4}[A-Z]{1}[A-Z\d]{1}Z[A-Z\d]{1})",
        text,
        re.IGNORECASE,
    )
    if m_gstin:
        fields["gstin_supplier"] = m_gstin.group(1).upper()

    fields["cgst"] = _parse_amount(_first_match(text, [r"cgst\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)"]))
    fields["sgst"] = _parse_amount(_first_match(text, [r"sgst\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)"]))
    fields["igst"] = _parse_amount(_first_match(text, [r"igst\s*[:\-]?\s*[₹]?\s*([\d,]+\.?\d*)"]))

    return fields


# ───────────────────── Common Fields (all types) ─────────────────

def _extract_common_fields(text: str) -> dict:
    """Extract PAN and GSTIN — present in almost every Indian financial document."""
    fields: dict = {}

    # PAN: 5 uppercase letters + 4 digits + 1 uppercase letter
    m_pan = re.search(r"\b([A-Z]{5}[0-9]{4}[A-Z]{1})\b", text)
    if m_pan:
        fields["pan"] = m_pan.group(1)

    # GSTIN: 2 digits + 5 uppercase + 4 digits + 1 uppercase + 1 alphanum + Z + 1 alphanum
    m_gstin = re.search(
        r"\b(\d{2}[A-Z]{5}\d{4}[A-Z]{1}[A-Z\d]{1}Z[A-Z\d]{1})\b",
        text,
    )
    if m_gstin:
        fields["gstin"] = m_gstin.group(1).upper()

    return fields


# ──────────────────────── Master extractor ───────────────────────

def extract_fields_heuristic(text: str, document_type: str) -> dict:
    """
    Extract structured fields from document text using regex.
    Returns a dict of {field: value}.  Null-valued keys are removed.
    """
    if not text:
        return {}

    fields: dict = {}

    if document_type == "salary_slip":
        fields.update(_extract_salary_slip_fields(text))
    elif document_type == "forex_document":
        fields.update(_extract_forex_fields(text))
    elif document_type == "fund_document":
        fields.update(_extract_fund_fields(text))
    elif document_type == "invoice":
        fields.update(_extract_invoice_fields(text))
    elif document_type in ("form_16", "itr_document"):
        # Form 16 / ITR share many salary-related fields
        fields.update(_extract_salary_slip_fields(text))

    # PAN + GSTIN extraction runs for all types
    fields.update(_extract_common_fields(text))

    # Strip None values
    return {k: v for k, v in fields.items() if v is not None}


# ═══════════════════════════════════════════════════════════════════
# PART 4 — LLM ENHANCEMENT (optional layer)
# ═══════════════════════════════════════════════════════════════════

# Fields that should ideally be populated per document type
_EXPECTED_FIELDS: dict = {
    "salary_slip":    ["gross_salary", "tds_deducted", "pf", "net_pay", "pan", "hra", "sec80c", "sec80d"],
    "form_16":        ["gross_salary", "tds_deducted", "pan", "hra", "sec80c", "sec80d"],
    "itr_document":   ["gross_salary", "tds_deducted", "pan"],
    "invoice":        ["invoice_number", "total_amount", "gstin_supplier"],
    "forex_document": ["currency_pair", "amount", "exchange_rate"],
    "fund_document":  ["fund_name", "nav", "units"],
}


def enhance_with_llm(text: str, heuristic_result: dict) -> dict:
    """
    Optionally improve heuristic results using an LLM (Ollama → OpenAI fallback).

    Rules:
      - If heuristic confidence > 0.7 and no expected fields are missing → skip LLM
      - If LLM unavailable → return heuristic_result unchanged with
        llm_enhancement = "unavailable"
      - If LLM JSON parse fails → return heuristic_result with
        llm_enhancement = "failed"
    """
    result = dict(heuristic_result)
    doc_type = result.get("document_type", "unknown")
    extracted = result.get("extracted_fields", {})

    # Always run the LLM if we have one to verify and correct the values, 
    # as OCR artifacts (like ₹ -> 2) might have corrupted the heuristic extractions.
    # We will pass both empty and filled fields for verification.
    
    try:
        from agents import get_llm
        from langchain_core.messages import HumanMessage, SystemMessage

        llm = get_llm()
        expected = _EXPECTED_FIELDS.get(doc_type, [])
        prompt = (
            f"Document text (first 2000 chars):\n{text[:2000]}\n\n"
            f"Heuristic classification: {doc_type}\n"
            f"Initial heuristic extraction: {json.dumps(extracted, indent=2)}\n"
            f"All expected fields for this doc type: {', '.join(expected)}\n\n"
            "Tasks:\n"
            "1. Confirm if document_type is correct. "
            "If wrong, provide the correct type from: "
            "salary_slip, form_16, itr_document, invoice, gst_return, "
            "forex_document, fund_document, tax_notice, unknown.\n"
            "2. Verify the initial heuristic extraction. "
            "CRITICAL OCR WARNING: The OCR engine heavily misreads the Indian Rupee symbol '₹' as the digit '2', or symbols like '#' or '%'. "
            "For example, '₹43,750' is often erroneously scanned as '243,750.00', '#1,31,250' or '%65,62'. "
            "You MUST mathematically analyze the totals to detect and strip these false leading '2's, '#'s, or '%'s. "
            "3. VERY IMPORTANT FOR SALARY SLIPS: Tax calculations require ANNUAL (Year-To-Date) figures. "
            "If the document is a salary slip, you MUST extract the YTD (Year To Date) values if available. "
            "If YTD is not available, you MUST explicitly multiply the monthly amounts by 12 and output the annualized totals for gross_salary, basic_pay, hra, sec80c, sec80d, net_pay, etc.\n"
            "4. Extract any missing expected fields.\n"
            "5. Return ONLY valid JSON containing ALL corrected and annualized fields — no explanation:\n"
            '{"document_type_confirmed": true, '
            '"corrected_type": null, '
            '"filled_fields": {"field_name": value}}'
        )


        messages = [
            SystemMessage(
                content=(
                    "You are a financial document parser specialising in Indian "
                    "financial documents. Extract data accurately. Return only valid JSON."
                )
            ),
            HumanMessage(content=prompt),
        ]

        response = llm.invoke(messages)
        content = response.content.strip()

        # Strip markdown code fences if present
        for prefix in ("```json", "```"):
            if content.startswith(prefix):
                content = content[len(prefix):]
                break
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()

        llm_data = json.loads(content)

        # Apply type correction
        if not llm_data.get("document_type_confirmed") and llm_data.get("corrected_type"):
            corrected = llm_data["corrected_type"]
            if corrected in _ALL_KNOWN_TYPES:
                result["document_type"] = corrected
                result["is_relevant_for_regime"] = corrected in _REGIME_RELEVANT
                result["is_relevant_for_forex"] = corrected in _FOREX_RELEVANT
                result["is_relevant_for_fund"] = corrected in _FUND_RELEVANT

        # Merge filled fields
        filled = llm_data.get("filled_fields") or {}
        if filled and isinstance(filled, dict):
            current = result.get("extracted_fields", {})
            current.update({k: v for k, v in filled.items() if v is not None})
            result["extracted_fields"] = current

        result["detected_by"] = "llm"
        result["llm_enhancement"] = "applied"

    except ImportError:
        result["llm_enhancement"] = "unavailable"
    except RuntimeError:
        # Both Ollama and OpenAI failed — normal if LLM is not configured
        result["llm_enhancement"] = "unavailable"
    except json.JSONDecodeError as exc:
        logger.debug("LLM returned non-JSON response: %s", exc)
        result["llm_enhancement"] = "failed"
    except Exception as exc:
        logger.debug("LLM enhancement failed: %s", exc)
        result["llm_enhancement"] = "failed"

    return result


# ═══════════════════════════════════════════════════════════════════
# PART 5 — MASTER PROCESS FUNCTION
# ═══════════════════════════════════════════════════════════════════

_EMPTY_RESULT: dict = {
    "document_type": "unknown",
    "confidence": 0.0,
    "detected_by": "heuristic",
    "is_relevant_for_regime": False,
    "is_relevant_for_forex": False,
    "is_relevant_for_fund": False,
    "extracted_text_preview": "",
    "extracted_fields": {},
    "llm_enhancement": "unavailable",
}


def process_document(file_path: str, filename: str) -> dict:
    """
    Master document processing function.

    Flow:
      1. Extract text   (OCR / PDF / DOCX)
      2. Classify       (heuristic keyword scoring)
      3. Extract fields (regex)
      4. LLM enhance    (optional — silently skipped if unavailable)

    NEVER raises an exception — all errors produce a safe fallback dict.
    """
    try:
        # Step 1: Text extraction
        text = extract_text_from_file(file_path)

        # Step 2: Classification
        classification = classify_document_heuristic(text, filename)

        # Step 3: Field extraction
        raw_fields = extract_fields_heuristic(text, classification["document_type"])

        # Step 4: Assemble intermediate result
        intermediate = {
            **classification,
            "extracted_text_preview": text[:500] if text else "",
            "extracted_fields": raw_fields,
            "llm_enhancement": "unavailable",
        }

        # Step 5: Optional LLM enhancement
        result = enhance_with_llm(text, intermediate)

        # Ensure preview always present
        result.setdefault("extracted_text_preview", text[:500] if text else "")

        return result

    except Exception as exc:
        logger.error("process_document failed for %s: %s", file_path, exc, exc_info=True)
        err = dict(_EMPTY_RESULT)
        err["error"] = str(exc)
        return err


# ═══════════════════════════════════════════════════════════════════
# BACKWARD-COMPAT WRAPPER — keeps existing document.py route working
# ═══════════════════════════════════════════════════════════════════

class DocumentAgent:
    """
    Thin wrapper around process_document() that returns the dict shape
    expected by api/routes/document.py.

    Existing call:
        result = DocumentAgent().handle(file_path)
    remains unchanged.
    """

    def handle(self, file_path: str) -> dict:
        filename = os.path.basename(file_path) if file_path else ""
        result = process_document(file_path, filename)

        fields = result.get("extracted_fields", {})
        doc_type = result.get("document_type", "unknown")
        is_relevant = result.get("is_relevant_for_regime", False)

        # Build a human-readable relevance reason
        if is_relevant:
            relevance_reason = None
        else:
            relevance_reason = (
                f"This {doc_type.replace('_', ' ')} document is not directly applicable "
                "for regime calculations. Please upload a salary slip or Form 16."
            )

        return {
            **fields,
            # ── Fields expected by ExtractedDataResponse schema ──
            "document_type":          doc_type,
            "is_relevant_for_regime": is_relevant,
            "relevance_reason":       relevance_reason,
            "gross_salary":           fields.get("gross_salary"),
            "tds_deducted":           fields.get("tds_deducted"),
            "pf":                     fields.get("pf") or fields.get("pf_deduction"),
            "pan":                    fields.get("pan"),
            "gstin":                  fields.get("gstin") or fields.get("gstin_supplier"),
            # ── Extended / new fields ──
            "confidence":             result.get("confidence", 0.0),
            "detected_by":            result.get("detected_by", "heuristic"),
            "llm_enhancement":        result.get("llm_enhancement", "unavailable"),
            "is_relevant_for_forex":  result.get("is_relevant_for_forex", False),
            "is_relevant_for_fund":   result.get("is_relevant_for_fund", False),
            "extracted_text_preview": result.get("extracted_text_preview", ""),
        }
